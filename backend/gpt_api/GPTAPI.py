from dotenv import load_dotenv
import base64
import os
import requests
import json
from pdf2image import convert_from_path
from fpdf import FPDF
import re
from .assistant_instructions import assistant_instructions
from docx import Document

import logging
import pickle
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.auth.exceptions import RefreshError

load_dotenv()

# Define current directory, for relative paths
current_dir = os.path.dirname(os.path.abspath(__file__))

# Make configurations for Google Docs

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Updated Scopes
# Constants
SCOPES = ['https://www.googleapis.com/auth/documents',
          'https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = os.path.join(current_dir, "credentials.json")
CREDENTIALS_PATH = os.path.join(current_dir, "credentials.json")
TOKEN_PICKLE_PATH = os.path.join(current_dir, "token.pickle")

# Function to encode the image
def encode_image(image_path):
  with open(image_path, "rb") as image_file:
    return base64.b64encode(image_file.read()).decode('utf-8')

# Class for GPT API instance
class GPTAPI:
    def __init__(self):
        self.API_KEY = os.getenv("API_KEY")
        self.MAX_TOKENS_PER_API_CALL = os.getenv("MAX_TOKENS_PER_API_CALL")
        self.gpt_model = "chatgpt-4o-latest"
        self.assistant_instructions = assistant_instructions

    def gpt_get_response(self, image_basename):
        print(f"Making request for OpenAI API...")

        # Getting the base64 string for image
        image_path = os.path.join(current_dir, 'output_0_areas', f'{image_basename}.png')
        base64_image = encode_image(image_path)

        # Define headers
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.API_KEY}"
        }

        # Define parameters for API call
        payload = {
            "model": self.gpt_model,
            "messages": [
                {
                "role": "user",
                "content": [
                    {
                    "type": "text",
                    "text": self.assistant_instructions
                    },
                    {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}"
                    }
                    }
                ]
                }
            ],
            "max_tokens": int(self.MAX_TOKENS_PER_API_CALL)
        }

        # Make API call
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)

        return response
    
    def generate_json(self, image_basename):
        # Make GPT API call for respective image
        response = self.gpt_get_response(image_basename)
        response_data = response.json()
        self.save_response_data_as_json(image_basename, response_data)

    def save_img(self, image_basename, image):
        image_path = os.path.join(current_dir, 'output_0_areas', image_basename + ".png")
        image.save(image_path, 'PNG')

    def save_response_data_as_json(self, image_basename, response_data) -> None:
        print(f"Saving data as JSON: {image_basename}")

        # Extract the 'choices' field from the response
        choices_data = response_data['choices']

        # Define the content string
        content = choices_data[0]['message']['content']

        # Clean content string
        content = content.replace("```", "")
        content = content.replace("json", "")
        content = content.replace("\n", "")

        # Transform content string to JSON
        try:
            # Attempt to load the string as JSON
            parsed_json = json.loads(content, strict=False)
        except json.JSONDecodeError as e:
            print(f"Error decoding JSON: {e}")

        # Specify the filename you want to write to
        filepath = os.path.join(current_dir, 'output_1_jsons', f'{image_basename}.json')
        
        # Write the 'choices' data to a file
        with open(filepath, 'w', encoding='utf-8') as file:
            json.dump(parsed_json, file, indent=2, ensure_ascii=False)

    def find_matching_files(self, file_basename, directory):
        """
        This function has O(N) time complexity, where N is the total number of files in the given directory.
        """
        # This regex pattern will match any file that contains the file_basename anywhere in the filename.
        pattern = re.compile(rf"{re.escape(file_basename)}")

        # List all files in the given directory
        files = os.listdir(directory)

        # Filter files that match the pattern (list their path containing directory)
        matching_files = [os.path.join(directory, file) for file in files if pattern.match(file)]

        return matching_files

    def insert_json_data_into_pdf_instance(self, json_data, pdf_instance, question_num):
        # Add page to PDF
        pdf_instance.add_page()

        # Display Question Number
        pdf_instance.set_font("DejaVu", style='B', size=12)
        pdf_instance.cell(200, 7, txt=f"Questão {question_num})", ln=True)
        pdf_instance.cell(200, 7, txt="", ln=True)

        # Display Question 'enunciado'
        question_enunciado = json_data["enunciado"]
        pdf_instance.set_font("DejaVu", style='', size=11)
        pdf_instance.multi_cell(190, 7, txt=question_enunciado)

        if json_data['tipo'] == "Objetiva":
            pdf_instance.cell(200, 5, txt="", ln=True)
            for key, value in json_data['resposta'].items():
                if key != "alternativaCorreta":
                    text_alternativa = key.lower() + ')  ' + value['alternativa']
                    pdf_instance.multi_cell(190, 7, txt=text_alternativa)
                    pdf_instance.cell(200, 4, txt="", ln=True)

        # Display Question 'resposta';
        pdf_instance.cell(190, 10, txt="", ln=True)
        pdf_instance.set_font("DejaVu", style='B', size=12)
        pdf_instance.cell(200, 7, txt="Solução:", ln=True)
        pdf_instance.cell(200, 7, txt="", ln=True)

        pdf_instance.set_font("DejaVu", style='', size=11)
        if json_data['tipo'] == "Discursiva":
            pdf_instance.multi_cell(190, 7, txt=json_data['resposta'])
        elif json_data['tipo'] == "Objetiva":
            for key, value in json_data['resposta'].items():
                if key != "alternativaCorreta":
                    textoExplicativo = key.lower() + ')  ' + value['textoExplicativo']
                    pdf_instance.multi_cell(180, 7, txt=textoExplicativo)
                    pdf_instance.cell(200, 4, txt="", ln=True)

            pdf_instance.cell(190, 10, txt="", ln=True)
            pdf_instance.set_font("DejaVu", style='B', size=12)
            pdf_instance.cell(200, 7, txt='Alternativa correta: ', ln=True)

            pdf_instance.set_font("DejaVu", style='', size=11)
            pdf_instance.cell(200, 7, txt=json_data['resposta']['alternativaCorreta'].lower(), ln=True)

    def generate_pdf_from_jsons(self, file_basename):

        # Define JSON files
        json_directory = os.path.join(current_dir, 'output_1_jsons')
        files_paths = self.find_matching_files(file_basename, json_directory)

        # Create PDF instance
        pdf_instance = FPDF()
        font_dejavu_path = os.path.join(current_dir, 'dejavu_font', 'ttf', 'DejaVuSans.ttf')
        pdf_instance.add_font('DejaVu', '', font_dejavu_path, uni=True)
        font_dejavu_bold_path = os.path.join(current_dir, 'dejavu_font', 'ttf', 'DejaVuSans-Bold.ttf')
        pdf_instance.add_font('DejaVu', 'B', font_dejavu_bold_path, uni=True)  

        # Iterate over JSON files, and add content do PDF
        for i, file_path in enumerate(files_paths):
            with open(file_path, 'r') as file:
                json_data = json.load(file)

            question_num = i+1
            self.insert_json_data_into_pdf_instance(json_data, pdf_instance, question_num)

        # Save PDF
        output_path_name = os.path.join(current_dir, 'output_2_pdfs', file_basename + "_resolvida.pdf")
        pdf_instance.output(output_path_name)

    def insert_json_data_into_doc_instance(self, json_data, doc_instance, question_num):
        # Display Question Number
        doc_instance.add_paragraph().add_run(f"Questão {question_num})").bold = True

        # Display Question 'enunciado'
        doc_instance.add_paragraph().add_run(json_data['enunciado'])

        if json_data['tipo'] == "Objetiva":
            for key, value in json_data['resposta'].items():
                if key != "alternativaCorreta":
                    p1 = doc_instance.add_paragraph()
                    p1.add_run(key.lower() + ')  ' + value['alternativa'])
                    
            doc_instance.add_paragraph('\n')

        # Display Question 'resposta';
        p2 = doc_instance.add_paragraph()
        p2.add_run('Solução:').bold = True

        if json_data['tipo'] == "Discursiva":
            doc_instance.add_paragraph().add_run(json_data['resposta'])
        elif json_data['tipo'] == "Objetiva":
            for key, value in json_data['resposta'].items():
                if key != "alternativaCorreta":
                    p3 = doc_instance.add_paragraph(key.lower() + ')  ')
                    p3.add_run(value['textoExplicativo'])

            doc_instance.add_paragraph('\n')

            p4 = doc_instance.add_paragraph()
            p4.add_run('Alternativa correta: ').bold = True

            doc_instance.add_paragraph(json_data['resposta']['alternativaCorreta'].lower())

    def generate_doc_from_jsons(self, file_basename):
        # Define JSON files
        json_directory = os.path.join(current_dir, 'output_1_jsons')
        files_paths = self.find_matching_files(file_basename, json_directory)

        # Create Doc instance
        doc_instance = Document()

        # Iterate over JSON files, and add content do PDF
        for i, file_path in enumerate(files_paths):
            with open(file_path, 'r') as file:
                json_data = json.load(file)

            question_num = i+1
            self.insert_json_data_into_doc_instance(json_data, doc_instance, question_num)

            # Add page break only if it is not the last one
            if i != len(files_paths) - 1:
                doc_instance.add_page_break()

        # Save Doc
        output_path_name = os.path.join(current_dir, 'output_2_pdfs', file_basename + "_resolvida.docx")
        doc_instance.save(output_path_name)

    def gpt_solver(self, pdf_filename, original_pdf_filename):
        """
        Endpoint to activate GPT API on previously uploaded file, after having confirmed the payment.
        """

        # Define the path of uploaded PDF file, based on the filename
        pdf_path = os.path.join(current_dir, "uploaded_files", pdf_filename)

        # Define file basename (without extension)
        file_basename, _ = os.path.splitext(pdf_filename)

        # Generate iterable of PDF pages as images
        pdf_pages_as_imgs = convert_from_path(pdf_path)

        # Iterate images
        for i, image in enumerate(pdf_pages_as_imgs):
            # Define image basename (without extension)
            image_basename = f"{file_basename}_page_{str(i+1).zfill(4)}" # Without extension

            # Save image as .png
            self.save_img(image_basename, image)

            # Call GPT API to generate JSON from processed image
            self.generate_json(image_basename)

        # Generate PDF from JSONs
        self.generate_pdf_from_jsons(file_basename)

        # Generate DOCX from JSONs
        self.generate_doc_from_jsons(file_basename)

        # Generate Google Docs from JSONs
        google_docs_url = ""
        google_docs_url = self.create_google_docs_with_jsons_contents(file_basename, [], original_pdf_filename)
        
        return google_docs_url
    
    # Google Docs methods

    def authenticate_google_services(self):
        """Authenticates the user and returns both Docs and Drive service objects."""
        creds = None
        token_path = TOKEN_PICKLE_PATH
        credentials_path = CREDENTIALS_PATH

        # Load existing credentials
        if os.path.exists(token_path):
            with open(token_path, 'rb') as token_file:
                creds = pickle.load(token_file)
                logger.info("Loaded credentials from token.pickle.")

        # If there are no valid credentials, let the user log in.
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    logger.info("Attempting to refresh expired credentials...")
                    creds.refresh(Request())
                    logger.info("Credentials refreshed successfully.")
                except RefreshError:
                    logger.error("Failed to refresh credentials. Token may have been revoked or expired.")
                    os.remove(token_path)
                    logger.info(f"Deleted invalid token file: {token_path}")
                    # Initiate a new authentication flow
                    flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
                    creds = flow.run_local_server(port=0)
                    logger.info("Authentication successful. New credentials obtained.")
            else:
                logger.info("No valid credentials available. Initiating authentication flow...")
                flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
                creds = flow.run_local_server(port=0)
                logger.info("Authentication successful. Credentials obtained.")

            # Save the credentials for the next run
            with open(token_path, 'wb') as token_file:
                pickle.dump(creds, token_file)
                logger.info(f"Saved new credentials to {token_path}.")

        # Build the Google Docs and Drive services
        docs_service = build('docs', 'v1', credentials=creds)
        drive_service = build('drive', 'v3', credentials=creds)
        logger.info("Built Google Docs and Drive service objects.")

        return docs_service, drive_service

    def create_google_document(self, docs_service, title="New Document"):
        """
        Creates a new Google Document and returns its ID.
        """
        document = docs_service.documents().create(body={'title': title}).execute()
        document_id = document.get('documentId')
        logger.info(f"Created document with ID: {document_id}")
        return document_id

    def insert_formatted_content(self, docs_service, document_id, content_list):
        """
        Inserts formatted JSON content into the Google Document.
        
        :param docs_service: Authenticated Google Docs service object.
        :param document_id: ID of the Google Document.
        :param content_list: List of content items with formatting directives.
        """
        requests = []
        current_index = 1  # Start after the beginning of the document

        for item in content_list:
            if item['type'] == 'heading':
                # Insert heading text
                requests.append({
                    'insertText': {
                        'location': {
                            'index': current_index
                        },
                        'text': item['text'] + '\n'
                    }
                })
                # Apply heading style
                requests.append({
                    'updateParagraphStyle': {
                        'range': {
                            'startIndex': current_index,
                            'endIndex': current_index + len(item['text']) + 1  # +1 for newline
                        },
                        'paragraphStyle': {
                            'namedStyleType': f"HEADING_{item['level']}"
                        },
                        'fields': 'namedStyleType'
                    }
                })
                # Apply text styling (font size, bold, italic)
                requests.append({
                    'updateTextStyle': {
                        'range': {
                            'startIndex': current_index,
                            'endIndex': current_index + len(item['text'])
                        },
                        'textStyle': {
                            'fontSize': {
                                'magnitude': item.get('font_size', 12),
                                'unit': 'PT'
                            },
                            'bold': item.get('bold', False),
                            'italic': item.get('italic', False)
                        },
                        'fields': 'fontSize,bold,italic'
                    }
                })
                current_index += len(item['text']) + 1  # Move index past the inserted text and newline

            elif item['type'] == 'paragraph':
                # Insert paragraph text
                requests.append({
                    'insertText': {
                        'location': {
                            'index': current_index
                        },
                        'text': item['text'] + '\n'
                    }
                })
                # Apply paragraph styling
                requests.append({
                    'updateParagraphStyle': {
                        'range': {
                            'startIndex': current_index,
                            'endIndex': current_index + len(item['text']) + 1
                        },
                        'paragraphStyle': {
                            'namedStyleType': 'NORMAL_TEXT'
                        },
                        'fields': 'namedStyleType'
                    }
                })
                # Apply text styling (font size, bold, italic)
                requests.append({
                    'updateTextStyle': {
                        'range': {
                            'startIndex': current_index,
                            'endIndex': current_index + len(item['text'])
                        },
                        'textStyle': {
                            'fontSize': {
                                'magnitude': item.get('font_size', 12),
                                'unit': 'PT'
                            },
                            'bold': item.get('bold', False),
                            'italic': item.get('italic', False)
                        },
                        'fields': 'fontSize,bold,italic'
                    }
                })
                current_index += len(item['text']) + 1

            elif item['type'] == 'page_break':
                # Insert a page break
                requests.append({
                    'insertPageBreak': {
                        'location': {
                            'index': current_index
                        }
                    }
                })
                current_index += 1  # Page break occupies one index

            # Add more types as needed (e.g., images, tables)

        # Execute the batch update
        try:
            result = docs_service.documents().batchUpdate(
                documentId=document_id, body={'requests': requests}).execute()
            logger.info(f"Inserted formatted content into document ID: {document_id}")
        except HttpError as error:
            logger.error(f"An error occurred: {error}")
            raise

    def share_document(self, drive_service, document_id, share_emails, make_public=False):
        """
        Shares the Google Document with specified emails and/or makes it public.

        :param drive_service: Authenticated Google Drive service object.
        :param document_id: ID of the Google Document.
        :param share_emails: List of email addresses to share the document with.
        :param make_public: Boolean indicating if the document should be public.
        """
        permissions = []

        if make_public:
            permissions.append({
                'type': 'anyone',
                'role': 'reader'
            })

        for email in share_emails:
            permissions.append({
                'type': 'user',
                'role': 'writer',
                'emailAddress': email
            })

        for permission in permissions:
            try:
                drive_service.permissions().create(
                    fileId=document_id,
                    body=permission,
                    fields='id'
                ).execute()
                if permission['type'] == 'anyone':
                    logger.info(f"Granted {permission.get('role')} access to anyone (public).")
                else:
                    logger.info(f"Granted {permission.get('role')} access to {permission.get('emailAddress')}.")
            except HttpError as error:
                logger.error(f"An error occurred while sharing: {error}")

    def create_google_docs(self, share_emails, json_content, title="Generated Document", make_public=True):
        """
        Creates a Google Document with formatted JSON content spread across multiple pages.

        :param share_emails: List of email addresses to share the document with.
        :param json_content: JSON content to insert into the document.
        :param title: Title of the Google Document.
        :param make_public: Boolean indicating if the document should be public.
        """
        try:
            # Authenticate and build service objects
            docs_service, drive_service = self.authenticate_google_services()

            # Create the document
            document_id = self.create_google_document(docs_service, title=title)

            # Parse JSON content (assuming it's a list)
            if isinstance(json_content, str):
                content_list = json.loads(json_content)
            elif isinstance(json_content, list):
                content_list = json_content
            else:
                raise ValueError("json_content must be a JSON string or a list of items.")

            # Insert formatted content into the document
            self.insert_formatted_content(docs_service, document_id, content_list)

            # Share the document
            self.share_document(drive_service, document_id, share_emails, make_public=make_public)

            # Get the document URL
            full_url = f"https://docs.google.com/document/d/{document_id}/edit"
            logger.info(f"Google Docs URL: {full_url}")

            # Return the document URL
            return full_url

        except Exception as e:
            logger.critical("A critical error occurred: %s", e)
            return None


    def list_formatted_jsons_contents(self, common_basename):
        jsons_filenames = self.find_matching_files(common_basename, os.path.join(current_dir, "output_1_jsons"))

        # jsons_enunciados_filenames = [s for s in jsons_filenames if "_resposta" not in s]
        # jsons_respostas_filenames = [s for s in jsons_filenames if "_resposta" in s]

        json_content = []

        # if (len(jsons_enunciados_filenames) != len(jsons_respostas_filenames)):
        #     raise Exception
        
        # for i in range(len(jsons_enunciados_filenames)):
        for i in range(len(jsons_filenames)):
            # json_enunciado_filename = jsons_enunciados_filenames[i]
            # json_enunciado_filepath = os.path.join(current_dir, "output_1_jsons", json_enunciado_filename)
            # json_resposta_filename = jsons_respostas_filenames[i]
            # json_resposta_filepath = os.path.join(current_dir, "output_1_jsons", json_resposta_filename)
            json_filename = jsons_filenames[i]

            with open(json_filename, 'r') as file:
                json_data = json.load(file)

            # with open(json_enunciado_filepath, 'r') as file:
            #     json_enunciado_data = json.load(file)

            # with open(json_resposta_filepath, 'r') as file:
            #     json_resposta_data = json.load(file)

            # Display Question Number
            json_content.append({
                    "type": "paragraph",
                    "text": f"Questão {i+1})",
                    "font_size": 12,
                    "bold": True
            })

            json_content.append({
                    "type": "paragraph",
                    "text": " ",
                    "font_size": 12,
            })

            # Display Question 'enunciado'
            json_content.append({
                    "type": "paragraph",
                    "text": json_data['enunciado'],
                    "font_size": 12,
            })

            if json_data['tipo'] == "Objetiva":
                json_content.append({
                    "type": "paragraph",
                    "text": " ",
                    "font_size": 12,
                })

                for key, value in json_data['resposta'].items():
                    if key != "alternativaCorreta":
                        json_content.append({
                                "type": "paragraph",
                                "text": key.upper() + ')  ' + value['alternativa'],
                                "font_size": 12,
                        })
                        
                json_content.append({
                        "type": "paragraph",
                        "text": " ",
                        "font_size": 12,
                })
            else:
                json_content.append({
                    "type": "paragraph",
                    "text": " ",
                    "font_size": 12,
                })

            # Display Question 'resposta';
            json_content.append({
                "type": "paragraph",
                "text": 'Solução:',
                "font_size": 12,
                "bold": True
            })

            json_content.append({
                "type": "paragraph",
                "text": " ",
                "font_size": 12,
            })

            if json_data['tipo'] == "Discursiva":
                json_content.append({
                    "type": "paragraph",
                    "text": json_data['resposta'],
                    "font_size": 12,
                })
                json_content.append({
                    "type": "paragraph",
                    "text": " ",
                    "font_size": 12,
                })
            elif json_data['tipo'] == "Objetiva":
                for key, value in json_data['resposta'].items():
                    if key != "alternativaCorreta":
                        json_content.append({
                            "type": "paragraph",
                            "text": key.upper() + ')  ' + value['textoExplicativo'],
                            "font_size": 12,
                        })
                        json_content.append({
                            "type": "paragraph",
                            "text": " ",
                            "font_size": 12,
                        })

                json_content.append({
                    "type": "paragraph",
                    "text": " ",
                    "font_size": 12,
                })

                json_content.append({
                    "type": "paragraph",
                    "text": 'Alternativa correta: ' + json_data['resposta']['alternativaCorreta'].upper(),
                    "font_size": 12,
                    "bold": True
                })

            # Break page if not the last question
            if i != len(jsons_filenames) - 1:
                json_content.append({
                    "type": "page_break"
                })
                    
        return json_content

    def create_google_docs_with_jsons_contents(self, common_basename, share_emails, original_pdf_filename):
        
        # JSON content
        json_content = self.list_formatted_jsons_contents(common_basename)  

        # Create the Google Doc
        document_url = self.create_google_docs(
            share_emails=share_emails,
            json_content=json_content,
            title=f'{original_pdf_filename.rsplit(".", 1)[0]} - Resolvida',
            make_public=True  # Set to True to make the document publicly accessible
        )

        # Return the document URL
        return document_url